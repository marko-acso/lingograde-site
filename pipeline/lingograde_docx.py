#!/usr/bin/env python3
"""
LingoGrade Report Generator v3.0
Reads JSON output from the LingoGrade Assessment Protocol and generates
a branded .docx report with side-by-side Assessed/Native language layout.
Supports any assessed language — German only appears when German IS the assessed language.

Usage:
    python lingograde_docx.py input.json [output.docx]

If output path is omitted, generates: LingoGrade_Report_<student>_<date>.docx
"""

import json
import sys
import os
from datetime import datetime, timedelta

# Import comprehensive translations
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
try:
    from lg_translations import TOC_FULL, TOC_HEADING, CONFIDENTIAL as CONF_MAP, PREPARED_FOR, CLOSING, MARKO_NATIVE as MARKO_MAP, EMAIL as EMAIL_TR
    HAS_TRANSLATIONS = True
except ImportError:
    HAS_TRANSLATIONS = False

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Logo Path (same directory as script) ──
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(SCRIPT_DIR, "logo.png")
HAS_LOGO = os.path.exists(LOGO_PATH)


# ── Brand Colors ──
NAVY = RGBColor(0x1A, 0x3A, 0x5C)
ACCENT = RGBColor(0x25, 0x63, 0xAB)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
INK = RGBColor(0x1C, 0x1C, 0x1C)
INK_SOFT = RGBColor(0x5A, 0x5A, 0x5A)
INK_LIGHT = RGBColor(0x8A, 0x8A, 0x8A)
SUCCESS = RGBColor(0x27, 0xAE, 0x60)
ERROR_C = RGBColor(0xC0, 0x39, 0x2B)
HEADER_BG = "1A3A5C"
ACCENT_BG = "2563AB"
LIGHT_BG = "F0F4F8"
WHITE_BG = "FFFFFF"
BORDER_C = "CCCCCC"

# ── Trilingual Translation Map ──
# Maps German terms to (English, Native) for common assessment terms
TRANSLATIONS = {
    # Languages
    "Russisch": {"bg": "Руски", "en": "Russian", "ru": "Русский", "uk": "Російська", "fr": "Russe", "es": "Ruso", "it": "Russo", "sr": "Руски", "hr": "Ruski"},
    "Englisch": {"bg": "Английски", "en": "English", "ru": "Английский", "uk": "Англійська", "fr": "Anglais", "es": "Inglés", "it": "Inglese", "sr": "Енглески", "hr": "Engleski"},
    "Französisch": {"bg": "Френски", "en": "French", "ru": "Французский", "uk": "Французька", "fr": "Français", "es": "Francés", "it": "Francese", "sr": "Француски", "hr": "Francuski"},
    "Spanisch": {"bg": "Испански", "en": "Spanish", "ru": "Испанский", "uk": "Іспанська", "fr": "Espagnol", "es": "Español", "it": "Spagnolo", "sr": "Шпански", "hr": "Španjolski"},
    "Italienisch": {"bg": "Италиански", "en": "Italian", "ru": "Итальянский", "uk": "Італійська", "fr": "Italien", "es": "Italiano", "it": "Italiano", "sr": "Италијански", "hr": "Talijanski"},
    "Serbisch": {"bg": "Сръбски", "en": "Serbian", "ru": "Сербский", "uk": "Сербська", "fr": "Serbe", "es": "Serbio", "it": "Serbo", "sr": "Српски", "hr": "Srpski"},
    "Kroatisch": {"bg": "Хърватски", "en": "Croatian", "ru": "Хорватский", "uk": "Хорватська", "fr": "Croate", "es": "Croata", "it": "Croato", "sr": "Хрватски", "hr": "Hrvatski"},
    "Ukrainisch": {"bg": "Украински", "en": "Ukrainian", "ru": "Украинский", "uk": "Українська", "fr": "Ukrainien", "es": "Ucraniano", "it": "Ucraino", "sr": "Украјински", "hr": "Ukrajinski"},
    "Bulgarisch": {"bg": "Български", "en": "Bulgarian", "ru": "Болгарский", "uk": "Болгарська", "fr": "Bulgare", "es": "Búlgaro", "it": "Bulgaro", "sr": "Бугарски", "hr": "Bugarski"},
    "Portugiesisch": {"bg": "Португалски", "en": "Portuguese", "ru": "Португальский", "uk": "Португальська", "fr": "Portugais", "es": "Portugués", "it": "Portoghese", "sr": "Португалски", "hr": "Portugalski"},
    "Polnisch": {"bg": "Полски", "en": "Polish", "ru": "Польский", "uk": "Польська", "fr": "Polonais", "es": "Polaco", "it": "Polacco", "sr": "Пољски", "hr": "Poljski"},
    "Deutsch": {"bg": "Немски", "en": "German", "ru": "Немецкий", "uk": "Німецька", "fr": "Allemand", "es": "Alemán", "it": "Tedesco", "sr": "Немачки", "hr": "Njemački"},
    # Error classes
    "Verbposition": {"ar": "موضع الفعل", "bg": "Позиция на глагола", "en": "Verb Position", "ru": "Позиция глагола", "uk": "Позиція дієслова", "fr": "Position du verbe", "es": "Posición del verbo", "it": "Posizione del verbo", "sr": "Позиција глагола", "hr": "Pozicija glagola"},
    "Kasus/Artikel": {"ar": "الحالات/أدوات التعريف", "bg": "Падежи/Членове", "en": "Cases/Articles", "ru": "Падежи/Артикли", "uk": "Відмінки/Артиклі", "fr": "Cas/Articles", "es": "Casos/Artículos", "it": "Casi/Articoli", "sr": "Падежи/Чланови", "hr": "Padeži/Članovi"},
    "Präpositionen": {"ar": "حروف الجر", "bg": "Предлози", "en": "Prepositions", "ru": "Предлоги", "uk": "Прийменники", "fr": "Prépositions", "es": "Preposiciones", "it": "Preposizioni", "sr": "Предлози", "hr": "Prijedlozi"},
    "Wortstellung": {"ar": "ترتيب الكلمات", "bg": "Словоред", "en": "Word Order", "ru": "Порядок слов", "uk": "Порядок слів", "fr": "Ordre des mots", "es": "Orden de palabras", "it": "Ordine delle parole", "sr": "Ред речи", "hr": "Red riječi"},
    "Zeitformen": {"ar": "الأزمنة", "bg": "Времена", "en": "Tenses", "ru": "Времена", "uk": "Часи", "fr": "Temps", "es": "Tiempos", "it": "Tempi", "sr": "Времена", "hr": "Vremena"},
    "Register/Wortwahl": {"ar": "السجل/اختيار الكلمات", "bg": "Регистър/Избор на думи", "en": "Register/Word Choice", "ru": "Регистр/Выбор слов", "uk": "Регістр/Вибір слів", "fr": "Registre/Choix de mots", "es": "Registro/Elección de palabras", "it": "Registro/Scelta lessicale", "sr": "Регистар/Избор речи", "hr": "Registar/Izbor riječi"},
    "Wortwahl": {"ar": "اختيار الكلمات", "bg": "Избор на думи", "en": "Word Choice", "ru": "Выбор слов", "uk": "Вибір слів", "fr": "Choix de mots", "es": "Elección de palabras", "it": "Scelta lessicale", "sr": "Избор речи", "hr": "Izbor riječi"},
    # Stability levels
    "niedrig": {"ar": "منخفض", "bg": "Нисък", "en": "Low", "ru": "Низкий", "uk": "Низький", "fr": "Faible", "es": "Bajo", "it": "Basso", "sr": "Низак", "hr": "Nizak"},
    "niedrig bis mittel": {"bg": "Нисък до среден", "en": "Low to Medium", "ru": "Низкий до среднего", "uk": "Низький до середнього", "fr": "Faible à moyen", "es": "Bajo a medio", "it": "Basso-medio", "sr": "Низак до средњи", "hr": "Nizak do srednji"},
    "mittel": {"ar": "متوسط", "bg": "Среден", "en": "Medium", "ru": "Средний", "uk": "Середній", "fr": "Moyen", "es": "Medio", "it": "Medio", "sr": "Средnji", "hr": "Srednji"},
    "mittel bis hoch": {"bg": "Среден до висок", "en": "Medium to High", "ru": "Средний до высокого", "uk": "Середній до високого", "fr": "Moyen à élevé", "es": "Medio a alto", "it": "Medio-alto", "sr": "Средњи до висок", "hr": "Srednji do visok"},
    "hoch": {"ar": "مرتفع", "bg": "Висок", "en": "High", "ru": "Высокий", "uk": "Високий", "fr": "Élevé", "es": "Alto", "it": "Alto", "sr": "Висок", "hr": "Visok"},
    "instabil": {"ar": "غير مستقر", "bg": "Нестабилен", "en": "Unstable", "ru": "Нестабильный", "uk": "Нестабільний", "fr": "Instable", "es": "Inestable", "it": "Instabile", "sr": "Нестабилан", "hr": "Nestabilan"},
    "teilstabil": {"ar": "مستقر جزئياً", "bg": "Частично стабилен", "en": "Partially Stable", "ru": "Частично стабильный", "uk": "Частково стабільний", "fr": "Partiellement stable", "es": "Parcialmente estable", "it": "Parzialmente stabile", "sr": "Делимично стабилан", "hr": "Djelomično stabilan"},
    "stabil": {"ar": "مستقر", "bg": "Стабилен", "en": "Stable", "ru": "Стабильный", "uk": "Стабільний", "fr": "Stable", "es": "Estable", "it": "Stabile", "sr": "Стабилан", "hr": "Stabilan"},
    # Labels
    "Ziel": {"ar": "الهدف", "bg": "Цел", "en": "Goal", "ru": "Цель", "uk": "Ціль", "fr": "Objectif", "es": "Objetivo", "it": "Obiettivo", "sr": "Циљ", "hr": "Cilj"},
    "Schritt": {"ar": "الخطوة", "bg": "Стъпка", "en": "Step", "ru": "Шаг", "uk": "Крок", "fr": "Étape", "es": "Paso", "it": "Passo", "sr": "Корак", "hr": "Korak"},
    "Fokus": {"ar": "التركيز", "bg": "Фокус", "en": "Focus", "ru": "Фокус", "uk": "Фокус", "fr": "Focus", "es": "Enfoque", "it": "Focus", "sr": "Фокус", "hr": "Fokus"},
    "Was du gut machst": {"ar": "ما تفعله جيداً", "bg": "Какво правиш добре", "en": "What you do well", "ru": "Что ты делаешь хорошо", "uk": "Що ти робиш добре", "fr": "Ce que tu fais bien", "es": "Lo que haces bien", "it": "Cosa fai bene", "sr": "Шта добро радиш", "hr": "Što dobro radiš"},
    "Beispiel": {"ar": "مثال", "bg": "Пример", "en": "Example", "ru": "Пример", "uk": "Приклад", "fr": "Exemple", "es": "Ejemplo", "it": "Esempio", "sr": "Пример", "hr": "Primjer"},
    # Block D labels
    "Top Problems": {"ar": "المشاكل الرئيسية:", "bg": "Основни проблеми:", "en": "Top Problems:", "ru": "Основные проблемы:", "uk": "Основні проблеми:", "fr": "Problèmes principaux:", "es": "Problemas principales:", "it": "Problemi principali:", "sr": "Главни проблеми:", "hr": "Glavni problemi:"},
    "Key Recommendation": {"ar": "التوصية الرئيسية:", "bg": "Ключова препоръка:", "en": "Key Recommendation:", "ru": "Ключевая рекомендация:", "uk": "Ключова рекомендація:", "fr": "Recommandation clé:", "es": "Recomendación clave:", "it": "Raccomandazione chiave:", "sr": "Кључна препорука:", "hr": "Ključna preporuka:"},
    "Next Steps": {"ar": "الخطوات التالية:", "bg": "Следващи стъпки:", "en": "Next Steps:", "ru": "Следующие шаги:", "uk": "Наступні кроки:", "fr": "Prochaines étapes:", "es": "Próximos pasos:", "it": "Prossimi passi:", "sr": "Следећи кораци:", "hr": "Sljedeći koraci:"},
}

# Language code mapping from native language name
LANG_CODES = {
    "Russisch": "ru", "Russian": "ru", "Русский": "ru",
    "Ukrainisch": "uk", "Ukrainian": "uk", "Українська": "uk",
    "Französisch": "fr", "French": "fr", "Français": "fr",
    "Spanisch": "es", "Spanish": "es", "Español": "es",
    "Italienisch": "it", "Italian": "it", "Italiano": "it",
    "Serbisch": "sr", "Serbian": "sr", "Српски": "sr",
    "Kroatisch": "hr", "Croatian": "hr", "Hrvatski": "hr",
    "Englisch": "en", "English": "en",
    "Bulgarisch": "bg", "Bulgarian": "bg",
    "Arabisch": "ar", "Arabic": "ar", "العربية": "ar",
    "Türkisch": "tr", "Turkish": "tr", "Türkçe": "tr",
    "Albanisch": "sq", "Albanian": "sq", "Shqip": "sq",
    "Rumänisch": "ro", "Romanian": "ro", "Română": "ro",
    "Ungarisch": "hu", "Hungarian": "hu", "Magyar": "hu",
    "Tschechisch": "cs", "Czech": "cs", "Čeština": "cs",
    "Norwegisch": "no", "Norwegian": "no",
    "Schwedisch": "sv", "Swedish": "sv",
    "Finnisch": "fi", "Finnish": "fi",
    "Dänisch": "da", "Danish": "da",
    "Niederländisch": "nl", "Dutch": "nl",
    "Hindi": "hi",
    "Polnisch": "pl", "Polish": "pl",
    "Portugiesisch": "pt", "Portuguese": "pt",
}

def translate_term(term_de, lang_code):
    """Translate a German term to the target language. Returns (english, native)."""
    term_lower = term_de.strip().lower()
    for key, translations in TRANSLATIONS.items():
        if key.lower() == term_lower:
            en = translations.get("en", term_de)
            native = translations.get(lang_code, en)
            return en, native
    return term_de, term_de

def trilingual(term_de, lang_code):
    """Return 'English / Deutsch / Native' string for a term."""
    en, native = translate_term(term_de, lang_code)
    if lang_code == "en":
        return f"{en} / {term_de}"
    return f"{en} / {term_de} / {native}"

def get_native_label(term_de, lang_code):
    """Return just the native translation of a German label."""
    _, native = translate_term(term_de, lang_code)
    return native


def calc_reassessment_date(assessment_date_str, reassessment_period):
    """Calculate the next re-assessment date on the nearest Tue/Thu/Sat.
    
    Args:
        assessment_date_str: 'YYYY-MM-DD' format
        reassessment_period: German string like '3 Monate', '6 Wochen', etc.
    
    Returns:
        (date_obj, formatted_str, cal_com_url) or (None, None, None) if can't calculate
    """
    period_map = {
        "3 Monate": timedelta(weeks=13),
        "6 Monate": timedelta(weeks=26),
        "6 Wochen": timedelta(weeks=6),
        "8 Wochen": timedelta(weeks=8),
        "12 Wochen": timedelta(weeks=12),
    }
    delta = period_map.get(reassessment_period)
    if not delta:
        return None, None, None
    
    try:
        assess_date = datetime.strptime(assessment_date_str, "%Y-%m-%d").date()
    except (ValueError, TypeError):
        return None, None, None
    
    target = assess_date + delta
    # Assessment days: Tue=1, Thu=3, Sat=5 (Monday=0)
    assessment_days = [1, 3, 5]  # Tue, Thu, Sat
    
    # Find nearest assessment day on or after target
    for offset in range(7):
        candidate = target + timedelta(days=offset)
        if candidate.weekday() in assessment_days:
            # Format: "Tuesday, June 23, 2026"
            day_names = {1: "Tuesday", 3: "Thursday", 5: "Saturday"}
            formatted = candidate.strftime(f"{day_names[candidate.weekday()]}, %B %d, %Y")
            # Cal.com URL with date parameter
            cal_url = f"https://cal.com/marko.check/full-assessment?date={candidate.isoformat()}"
            return candidate, formatted, cal_url
    
    return None, None, None


def add_hyperlink(paragraph, url, text, color=None, bold=False, size=None):
    """Add a clickable hyperlink to a paragraph in python-docx."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    
    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" {nsdecls("r")}>'
        f'</w:hyperlink>'
    )
    
    new_run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr>'
        f'<w:rStyle w:val="Hyperlink"/>'
        f'<w:color w:val="{str(color) if color else "2563AB"}"/>'
        f'<w:u w:val="single"/>'
        f'{"<w:b/>" if bold else ""}'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    
    if size:
        sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="{int(size * 2)}"/>')
        new_run.find(qn('w:rPr')).append(sz)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar_existing = tcPr.find(qn('w:tcMar'))
    if tcMar_existing is not None:
        tcPr.remove(tcMar_existing)
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set individual borders on a cell. Each param is a dict like {'sz': '4', 'color': 'CCCCCC', 'val': 'single'}"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    for edge, data in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if data:
            elem = tcBorders.find(qn(f'w:{edge}'))
            if elem is not None:
                tcBorders.remove(elem)
            border_xml = f'<w:{edge} {nsdecls("w")} w:val="{data.get("val","single")}" w:sz="{data.get("sz","4")}" w:space="0" w:color="{data.get("color","CCCCCC")}"/>'
            tcBorders.append(parse_xml(border_xml))


def set_table_borders(tbl, color="CCCCCC", sz="4"):
    """Apply light borders to all cells in a table."""
    border = {'sz': sz, 'color': color, 'val': 'single'}
    for row in tbl.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, left=border, right=border)


def add_paragraph(cell, text, bold=False, size=9, color=INK, space_after=Pt(2), space_before=Pt(0), alignment=None):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    return p


def create_report(data, output_path=None):
    meta = data["metadata"]
    sections = {s["id"]: s for s in data["sections"]}
    cefr = data.get("cefr_footer", {})
    block_d = data.get("block_d", {})

    student = meta["student_name"]
    display = meta.get("display_name", student)
    native_display = meta.get("native_name", display)
    date = meta["date"]
    package = meta["package"]
    native_lang = meta["native_language"]
    package_label = "Full Assessment" if package == "full" else "Quick Assessment"
    lang_code = LANG_CODES.get(native_lang, "en")
    assessed_lang = meta.get("assessed_language", "Deutsch")
    assessed_code = LANG_CODES.get(assessed_lang, "de")
    assessed_en, _ = translate_term(assessed_lang, "en")

    # Calculate re-assessment date and booking link
    reassess_period = meta.get("recommended_reassessment", "")
    reassess_date_obj, reassess_date_str, reassess_cal_url = calc_reassessment_date(date, reassess_period)

    if not output_path:
        safe_name = student.replace(" ", "_").replace(".", "")
        output_path = f"LingoGrade_Report_{safe_name}_{date}.docx"

    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # ── Default Font ──
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(9)
    style.font.color.rgb = INK

    content_width = Cm(17.4)  # A4 minus margins
    col_width = Cm(8.5)

    # ══════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════

    # ── Navy header bar: Logo left, student info right ──
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Logo cell
    lc = header_table.cell(0, 0)
    set_cell_shading(lc, HEADER_BG)
    set_cell_margins(lc, top=120, bottom=120, left=160, right=80)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if HAS_LOGO:
        run = lp.add_run()
        run.add_picture(LOGO_PATH, height=Cm(1.6))
    else:
        run = lp.add_run("Lingo")
        run.font.name = 'Arial'; run.font.size = Pt(20); run.font.color.rgb = WHITE; run.font.bold = True
        run1b = lp.add_run("Grade")
        run1b.font.name = 'Arial'; run1b.font.size = Pt(20); run1b.font.color.rgb = RGBColor(0x93, 0xB5, 0xD0); run1b.font.bold = True
    # Student info cell
    rc = header_table.cell(0, 1)
    set_cell_shading(rc, HEADER_BG)
    set_cell_margins(rc, top=120, bottom=120, left=80, right=160)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = rp.add_run(student)
    run.font.name = 'Arial'; run.font.size = Pt(11); run.font.color.rgb = WHITE; run.font.bold = True
    rp2 = rc.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = rp2.add_run(date)
    run2.font.name = 'Arial'; run2.font.size = Pt(8); run2.font.color.rgb = RGBColor(0x93, 0xB5, 0xD0)
    rp3 = rc.add_paragraph()
    rp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run3 = rp3.add_run(f"{assessed_en} \u2022 {package_label}")
    run3.font.name = 'Arial'; run3.font.size = Pt(8); run3.font.color.rgb = RGBColor(0x93, 0xB5, 0xD0)

    # ── Title ──
    doc.add_paragraph("")

    # "Prepared exclusively for" — tactical empathy, personal ownership
    prepared_map = {
        "ru": f"Подготовлено эксклюзивно для {native_display}",
        "uk": f"Підготовлено ексклюзивно для {native_display}",
        "fr": f"Préparé exclusivement pour {native_display}",
        "es": f"Preparado exclusivamente para {native_display}",
        "it": f"Preparato esclusivamente per {native_display}",
        "sr": f"Припремљено ексклузивно за {native_display}",
        "hr": f"Pripremljeno ekskluzivno za {native_display}.",
        "bg": f"Подготвено ексклузивно за {native_display}",
        "ar": f"أُعد حصرياً لـ {native_display}",
        "pl": f"Przygotowane specjalnie dla {native_display}",
        "tr": f"{native_display} için özel olarak hazırlandı",
        "sq": f"Përgatitur ekskluzivisht për {native_display}",
        "ro": f"Pregătit exclusiv pentru {native_display}",
        "hu": f"Kizárólag {native_display} számára készítve",
    }
    prepared_text = prepared_map.get(lang_code, f"Prepared exclusively for {display}")
    p_excl = doc.add_paragraph()
    p_excl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_excl = p_excl.add_run(prepared_text)
    run_excl.font.name = 'Arial'; run_excl.font.size = Pt(10); run_excl.font.color.rgb = INK_LIGHT
    run_excl.font.italic = True
    p_excl.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Language Assessment Report")
    run.font.name = 'Arial'; run.font.size = Pt(20); run.font.color.rgb = INK; run.font.bold = True
    p.paragraph_format.space_after = Pt(2)

    duration = "25 min Conversation" if package == "full" else "15 min Conversation"
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"{date} \u2022 {assessed_en} \u2022 {duration} \u2022 {package_label}")
    run2.font.name = 'Arial'; run2.font.size = Pt(9); run2.font.color.rgb = INK_SOFT
    p2.paragraph_format.space_after = Pt(2)

    # Native language line
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    en_name_display, _ = translate_term(native_lang, lang_code)
    run3b = p3.add_run(f"Native Language: {en_name_display}")
    run3b.font.name = 'Arial'; run3b.font.size = Pt(9); run3b.font.color.rgb = INK_LIGHT
    p3.paragraph_format.space_after = Pt(4)

    # Confidential marker
    conf_map = {
        "ru": "КОНФИДЕНЦИАЛЬНО",
        "uk": "КОНФІДЕНЦІЙНО",
        "fr": "CONFIDENTIEL",
        "es": "CONFIDENCIAL",
        "it": "RISERVATO",
        "sr": "ПОВЕРЉИВО",
        "hr": "POVJERLJIVO",
        "bg": "ПОВЕРИТЕЛНО",
        "ar": "سري",
        "pl": "POUFNE",
        "tr": "GİZLİ",
        "sq": "KONFIDENCIALE",
        "ro": "CONFIDENȚIAL",
        "hu": "BIZALMAS",
    }
    p_conf = doc.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_conf = p_conf.add_run(conf_map.get(lang_code, "CONFIDENTIAL"))
    run_conf.font.name = 'Arial'; run_conf.font.size = Pt(7); run_conf.font.color.rgb = ERROR_C
    run_conf.font.bold = True; run_conf.font.all_caps = True
    p_conf.paragraph_format.space_after = Pt(8)

    # ── CEFR Score Cards (Active, Passive, Overall, Confidence) ──
    active = meta.get("active_cefr", "")
    passive = meta.get("passive_cefr", "")
    overall = active.split(".")[0] if "." in active else active
    confidence = meta.get("confidence", "")

    # Bilingual score labels: Assessed Language + Native Language
    score_labels_all = {
        "en": ["ACTIVE", "PASSIVE", "OVERALL", "CONFIDENCE"],
        "de": ["AKTIV", "PASSIV", "GESAMT", "SELBSTVERTRAUEN"],
        "ru": ["АКТИВНЫЙ", "ПАССИВНЫЙ", "ОБЩИЙ", "УВЕРЕННОСТЬ"],
        "uk": ["АКТИВНИЙ", "ПАСИВНИЙ", "ЗАГАЛЬНИЙ", "ВПЕВНЕНІСТЬ"],
        "fr": ["ACTIF", "PASSIF", "GLOBAL", "CONFIANCE"],
        "es": ["ACTIVO", "PASIVO", "GENERAL", "CONFIANZA"],
        "it": ["ATTIVO", "PASSIVO", "GLOBALE", "FIDUCIA"],
        "sr": ["АКТИВАН", "ПАСИВАН", "УКУПНО", "ПОУЗДАЊЕ"],
        "hr": ["AKTIVNO", "PASIVNO", "UKUPNO", "POUZDANJE"],
        "ar": ["نشط", "سلبي", "إجمالي", "ثقة"],
        "hi": ["सक्रिय", "निष्क्रिय", "समग्र", "आत्मविश्वास"],
        "pl": ["AKTYWNY", "PASYWNY", "OGÓLNY", "PEWNOŚĆ"],
        "pt": ["ATIVO", "PASSIVO", "GERAL", "CONFIANÇA"],
        "cs": ["AKTIVNÍ", "PASIVNÍ", "CELKOVÝ", "DŮVĚRA"],
        "hu": ["AKTÍV", "PASSZÍV", "ÖSSZESÍTETT", "MAGABIZTOSSÁG"],
        "no": ["AKTIV", "PASSIV", "SAMLET", "SELVTILLIT"],
        "sv": ["AKTIV", "PASSIV", "ÖVERGRIPANDE", "SJÄLVFÖRTROENDE"],
        "fi": ["AKTIIVINEN", "PASSIIVINEN", "KOKONAIS", "ITSELUOTTAMUS"],
        "da": ["AKTIV", "PASSIV", "SAMLET", "SELVTILLID"],
        "nl": ["ACTIEF", "PASSIEF", "ALGEMEEN", "ZELFVERTROUWEN"],
        "tr": ["AKTİF", "PASİF", "GENEL", "GÜVEN"],
        "sq": ["AKTIV", "PASIV", "PËRGJITHSHËM", "BESIM"],
        "ro": ["ACTIV", "PASIV", "GENERAL", "ÎNCREDERE"],
        "bg": ["АКТИВЕН", "ПАСИВЕН", "ОБЩ", "УВЕРЕНОСТ"],
    }
    score_labels_assessed = score_labels_all.get(assessed_code, score_labels_all.get("en"))
    score_labels_native = score_labels_all.get(lang_code, score_labels_all.get("en"))

    score_table = doc.add_table(rows=2, cols=4)
    score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = [active, passive, overall, confidence]
    colors = [ACCENT, ACCENT, INK, SUCCESS]

    for i, (val, clr) in enumerate(zip(values, colors)):
        # Label row — stacked trilingual
        c = score_table.cell(0, i)
        set_cell_margins(c, top=40, bottom=20, left=40, right=40)
        cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if assessed_code == lang_code:
            label_text = score_labels_assessed[i]
        else:
            label_text = f"{score_labels_assessed[i]} / {score_labels_native[i]}"
        r = cp.add_run(label_text)
        r.font.name = 'Arial'; r.font.size = Pt(6); r.font.color.rgb = INK_LIGHT; r.font.bold = True
        # Value row
        c2 = score_table.cell(1, i)
        set_cell_margins(c2, top=20, bottom=40, left=40, right=40)
        cp2 = c2.paragraphs[0]; cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = cp2.add_run(val)
        r2.font.name = 'Arial'; r2.font.size = Pt(18); r2.font.color.rgb = clr; r2.font.bold = True
    set_table_borders(score_table)

    # ── Detail strip (Error Focus, Stability, Self-Correction, Re-Assess) ──
    if package == "full":
        detail_table = doc.add_table(rows=2, cols=4)
        detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Bilingual detail strip labels: Assessed Language + Native Language
        d_labels_all = {
            "en": ["ERROR FOCUS", "STABILITY", "SELF-CORRECTION", "RE-ASSESS"],
            "de": ["FEHLERFOKUS", "STABILITÄT", "SELBSTKORREKTUR", "NEUBEWERTUNG"],
            "ru": ["ФОКУС ОШИБОК", "СТАБИЛЬНОСТЬ", "САМОКОРРЕКЦИЯ", "ПЕРЕОЦЕНКА"],
            "uk": ["ФОКУС ПОМИЛОК", "СТАБІЛЬНІСТЬ", "САМОКОРЕКЦІЯ", "ПЕРЕОЦІНКА"],
            "fr": ["ERREURS CLÉS", "STABILITÉ", "AUTO-CORRECTION", "RÉ-ÉVALUATION"],
            "es": ["ERRORES CLAVE", "ESTABILIDAD", "AUTOCORRECCIÓN", "REEVALUACIÓN"],
            "it": ["ERRORI CHIAVE", "STABILITÀ", "AUTOCORREZIONE", "RIVALUTAZIONE"],
            "sr": ["ФОКУС ГРЕШАКА", "СТАБИЛНОСТ", "САМОКОРЕКЦИЈА", "ПОНОВНА ПРОЦЕНА"],
            "hr": ["FOKUS GREŠAKA", "STABILNOST", "SAMOKOREKCIJA", "PONOVNA PROCJENA"],
            "ar": ["تركيز الأخطاء", "الاستقرار", "التصحيح الذاتي", "إعادة التقييم"],
            "pl": ["FOKUS BŁĘDÓW", "STABILNOŚĆ", "AUTOKOREKTA", "PONOWNA OCENA"],
            "pt": ["FOCO DE ERROS", "ESTABILIDADE", "AUTOCORREÇÃO", "REAVALIAÇÃO"],
            "cs": ["ZAMĚŘENÍ CHYB", "STABILITA", "AUTOKOREKCE", "PŘEHODNOCENÍ"],
            "hu": ["HIBAFÓKUSZ", "STABILITÁS", "ÖNJAVÍTÁS", "ÚJRAÉRTÉKELÉS"],
            "tr": ["HATA ODAĞI", "KARARLILIK", "ÖZ DÜZELTME", "YENİDEN DEĞERLENDİRME"],
            "sq": ["FOKUSI I GABIMEVE", "STABILITETI", "VETËKORRIGJIMI", "RIVLERËSIMI"],
            "ro": ["FOCUS ERORI", "STABILITATE", "AUTOCORECTARE", "REEVALUARE"],
            "bg": ["ФОКУС НА ГРЕШКИ", "СТАБИЛНОСТ", "САМОКОРЕКЦИЯ", "ПРЕОЦЕНКА"],
        }
        d_labels_assessed = d_labels_all.get(assessed_code, d_labels_all.get("en"))
        d_labels_native = d_labels_all.get(lang_code, d_labels_all.get("en"))

        # Translate error focus terms — stacked in one cell
        error_focus_raw = meta.get("error_focus", [])
        ef_lines = []
        for ef in error_focus_raw:
            if isinstance(ef, dict):
                assessed_val = ef.get(assessed_code, ef.get('de', ef.get('en', '')))
                native_val = ef.get(lang_code, ef.get('native', ''))
                if assessed_code == lang_code:
                    ef_lines.append(assessed_val)
                else:
                    ef_lines.append(f"{assessed_val} / {native_val}")
            else:
                assessed_term, _ = translate_term(ef, assessed_code)
                _, native_term = translate_term(ef, lang_code)
                # For German assessed, the original term IS the assessed term
                if assessed_code == "de":
                    assessed_term = ef
                if assessed_code == lang_code:
                    ef_lines.append(assessed_term)
                else:
                    ef_lines.append(f"{assessed_term} / {native_term}")

        # Translate stability and self-correction
        stab_de = meta.get("structural_stability", "")
        stab_assessed, _ = translate_term(stab_de, assessed_code)
        if assessed_code == "de": stab_assessed = stab_de
        _, stab_native = translate_term(stab_de, lang_code)
        if assessed_code == lang_code:
            stab_display = stab_assessed
        else:
            stab_display = f"{stab_assessed} / {stab_native}"

        corr_de = meta.get("self_correction", "")
        corr_assessed, _ = translate_term(corr_de, assessed_code)
        if assessed_code == "de": corr_assessed = corr_de
        _, corr_native = translate_term(corr_de, lang_code)
        if assessed_code == lang_code:
            corr_display = corr_assessed
        else:
            corr_display = f"{corr_assessed} / {corr_native}"

        # Re-assess translation
        reassess_de = meta.get("recommended_reassessment", "")
        reassess_map = {"3 Monate": "3 Months", "6 Monate": "6 Months", "6 Wochen": "6 Weeks", "8 Wochen": "8 Weeks", "12 Wochen": "12 Weeks"}
        reassess_native_map = {
            "ru": {"3 Monate": "3 месяца", "6 Monate": "6 месяцев", "6 Wochen": "6 недель", "8 Wochen": "8 недель", "12 Wochen": "12 недель"},
            "uk": {"3 Monate": "3 місяці", "6 Monate": "6 місяців", "6 Wochen": "6 тижнів", "8 Wochen": "8 тижнів", "12 Wochen": "12 тижнів"},
            "fr": {"3 Monate": "3 mois", "6 Monate": "6 mois", "6 Wochen": "6 semaines", "8 Wochen": "8 semaines", "12 Wochen": "12 semaines"},
            "es": {"3 Monate": "3 meses", "6 Monate": "6 meses", "6 Wochen": "6 semanas", "8 Wochen": "8 semanas", "12 Wochen": "12 semanas"},
            "it": {"3 Monate": "3 mesi", "6 Monate": "6 mesi", "6 Wochen": "6 settimane", "8 Wochen": "8 settimane", "12 Wochen": "12 settimane"},
            "sr": {"3 Monate": "3 месеца", "6 Monate": "6 месеци", "6 Wochen": "6 недеља", "8 Wochen": "8 недеља", "12 Wochen": "12 недеља"},
            "hr": {"3 Monate": "3 mjeseca", "6 Monate": "6 mjeseci", "6 Wochen": "6 tjedana", "8 Wochen": "8 tjedana", "12 Wochen": "12 tjedana"},
            "ar": {"3 Monate": "3 أشهر", "6 Monate": "6 أشهر", "6 Wochen": "6 أسابيع", "8 Wochen": "8 أسابيع", "12 Wochen": "12 أسبوع"},
            "pl": {"3 Monate": "3 miesiące", "6 Monate": "6 miesięcy", "6 Wochen": "6 tygodni", "8 Wochen": "8 tygodni", "12 Wochen": "12 tygodni"},
            "tr": {"3 Monate": "3 ay", "6 Monate": "6 ay", "6 Wochen": "6 hafta", "8 Wochen": "8 hafta", "12 Wochen": "12 hafta"},
            "sq": {"3 Monate": "3 muaj", "6 Monate": "6 muaj", "6 Wochen": "6 javë", "8 Wochen": "8 javë", "12 Wochen": "12 javë"},
            "ro": {"3 Monate": "3 luni", "6 Monate": "6 luni", "6 Wochen": "6 săptămâni", "8 Wochen": "8 săptămâni", "12 Wochen": "12 săptămâni"},
            "hu": {"3 Monate": "3 hónap", "6 Monate": "6 hónap", "6 Wochen": "6 hét", "8 Wochen": "8 hét", "12 Wochen": "12 hét"},
            "bg": {"3 Monate": "3 месеца", "6 Monate": "6 месеца", "6 Wochen": "6 седмици", "8 Wochen": "8 седмици", "12 Wochen": "12 седмици"},
        }
        reassess_en = reassess_map.get(reassess_de, reassess_de)
        reassess_assessed = reassess_native_map.get(assessed_code, {}).get(reassess_de, reassess_en)
        if assessed_code == "de": reassess_assessed = reassess_de
        elif assessed_code == "en": reassess_assessed = reassess_en
        reassess_native = reassess_native_map.get(lang_code, {}).get(reassess_de, reassess_en)
        if assessed_code == lang_code:
            reassess_display = reassess_assessed
        else:
            reassess_display = f"{reassess_assessed} / {reassess_native}"

        d_values = [
            "\n".join(ef_lines),
            stab_display,
            corr_display,
            reassess_display
        ]
        stability_color = RGBColor(0xE6, 0x7E, 0x22) if stab_de.lower() == "mittel" else SUCCESS
        correction_color = RGBColor(0xE6, 0x7E, 0x22) if corr_de.lower() == "mittel" else SUCCESS
        d_colors = [INK, stability_color, correction_color, INK_SOFT]

        for i in range(4):
            # Header — trilingual stacked
            h = detail_table.cell(0, i)
            set_cell_shading(h, "2C3E50")
            set_cell_margins(h, top=30, bottom=30, left=40, right=40)
            hp = h.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if assessed_code == lang_code:
                header_text = d_labels_assessed[i]
            else:
                header_text = f"{d_labels_assessed[i]} / {d_labels_native[i]}"
            hr = hp.add_run(header_text)
            hr.font.name = 'Arial'; hr.font.size = Pt(5); hr.font.color.rgb = WHITE; hr.font.bold = True
            # Value
            v = detail_table.cell(1, i)
            set_cell_margins(v, top=30, bottom=30, left=40, right=40)
            vp = v.paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vr = vp.add_run(d_values[i])
            vr.font.name = 'Arial'; vr.font.size = Pt(7); vr.font.color.rgb = d_colors[i]; vr.font.bold = True
        set_table_borders(detail_table, color="2C3E50")

    # ── "What This Report Contains" — value framing ──
    doc.add_paragraph("")
    
    # Use imported translations if available, else fallback
    if HAS_TRANSLATIONS:
        assessed_code = LANG_CODES.get(assessed_lang, "de")
        toc_heading_left = TOC_HEADING.get(assessed_code, TOC_HEADING.get("de", "Was dieser Bericht enthält"))
        toc_heading_native = TOC_HEADING.get(lang_code, TOC_HEADING.get("en", "What This Report Contains"))
        toc_de = TOC_FULL.get(assessed_code, TOC_FULL.get("de", []))
        toc_native = TOC_FULL.get(lang_code, TOC_FULL.get("en", []))
    else:
        toc_heading_left = "Was dieser Bericht enthält"
        toc_heading_native = "What This Report Contains"
        toc_de = ["Präzise Wahrnehmung", "Strukturelle Stärken", "Zentrale Erkenntnis", "Top 5 Probleme", "5 Hebel-Korrekturen", "Top 5 Lösungen", "Satzmotor", "Hausaufgaben A, B, C", "8-Wochen-Plan", "Zusammenfassung"]
        toc_native = toc_de

    if package != "full":
        # Quick package: fewer items
        toc_de = [toc_de[i] for i in [0, 2, 3, 4, 5, 6, 7, 8, 9] if i < len(toc_de)]
        toc_native = [toc_native[i] for i in [0, 2, 3, 4, 5, 6, 7, 8, 9] if i < len(toc_native)]

    # Build two-column TOC table
    toc_table = doc.add_table(rows=1 + len(toc_de), cols=2)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for idx, heading in enumerate([toc_heading_left, toc_heading_native]):
        cell = toc_table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BG)
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        p = cell.paragraphs[0]
        run = p.add_run(heading)
        run.font.name = 'Arial'; run.font.size = Pt(9); run.font.color.rgb = ACCENT; run.font.bold = True

    # Content rows
    for i in range(len(toc_de)):
        native_item = toc_native[i] if i < len(toc_native) else ""
        for idx, text in enumerate([toc_de[i], native_item]):
            cell = toc_table.cell(i + 1, idx)
            set_cell_margins(cell, top=20, bottom=20, left=100, right=100)
            if i % 2 == 0:
                set_cell_shading(cell, LIGHT_BG)
            p = cell.paragraphs[0]
            run = p.add_run(f"\u2022  {text}")
            run.font.name = 'Arial'; run.font.size = Pt(7); run.font.color.rgb = INK_SOFT

    set_table_borders(toc_table)
    doc.add_paragraph("")

    # ══════════════════════════════════════
    # HELPER: Full-width section banner
    # ══════════════════════════════════════
    # Banner labels: key → (german, {lang_code: native})
    BANNERS = {
        "problems": {
            "ar": "أهم {n} مشاكل هيكلية",
            "bg": "ТОП-{n} СТРУКТУРНИ ПРОБЛЕМИ",
            "de": "TOP {n} STRUKTURELLE PROBLEME",
            "ru": "ТОП-{n} СТРУКТУРНЫХ ПРОБЛЕМ",
            "uk": "ТОП-{n} СТРУКТУРНИХ ПРОБЛЕМ",
            "fr": "TOP {n} PROBLÈMES STRUCTURELS",
            "es": "TOP {n} PROBLEMAS ESTRUCTURALES",
            "it": "TOP {n} PROBLEMI STRUTTURALI",
            "sr": "ТОП {n} СТРУКТУРНИХ ПРОБЛЕМА",
            "hr": "TOP {n} STRUKTURNIH PROBLEMA",
            "en": "TOP {n} STRUCTURAL PROBLEMS",
        },
        "corrections": {
            "ar": "أفضل {n} تصحيحات رافعة",
            "bg": "ТОП-{n} КОРЕКЦИИ НА ЛОСТА",
            "de": "DIE {n} BESTEN HEBELKORREKTUREN",
            "ru": "ТОП-{n} КОРРЕКЦИЙ РЫЧАГА",
            "uk": "ТОП-{n} ВАЖІЛЬНИХ КОРЕКЦІЙ",
            "fr": "LES {n} MEILLEURES CORRECTIONS LEVIERS",
            "es": "LAS {n} MEJORES CORRECCIONES PALANCA",
            "it": "LE {n} MIGLIORI CORREZIONI LEVA",
            "sr": "НАЈБОЉИХ {n} КОРЕКЦИЈА ПОЛУГЕ",
            "hr": "NAJBOLJIH {n} KOREKCIJA POLUGE",
            "en": "TOP {n} LEVER CORRECTIONS",
        },
        "solutions": {
            "ar": "أفضل {n} حلول",
            "bg": "ТОП-{n} РЕШЕНИЯ",
            "de": "TOP {n} LÖSUNGEN",
            "ru": "ТОП-{n} РЕШЕНИЙ",
            "uk": "ТОП-{n} РІШЕНЬ",
            "fr": "TOP {n} SOLUTIONS",
            "es": "TOP {n} SOLUCIONES",
            "it": "TOP {n} SOLUZIONI",
            "sr": "ТОП {n} РЕШЕЊА",
            "hr": "TOP {n} RJEŠENJA",
            "en": "TOP {n} SOLUTIONS",
        },
        "engine": {
            "ar": "محرك الجمل",
            "bg": "РЕЧЕВИ ДВИГАТЕЛ",
            "de": "SATZMOTOR",
            "ru": "РЕЧЕВОЙ ДВИГАТЕЛЬ",
            "uk": "МОВНИЙ ДВИГУН",
            "fr": "MOTEUR DE PHRASE",
            "es": "MOTOR DE ORACIÓN",
            "it": "MOTORE DI FRASE",
            "sr": "МОТОР РЕЧЕНИЦЕ",
            "hr": "MOTOR REČENICE",
            "en": "SENTENCE ENGINE",
        },
        "activation": {
            "ar": "التنشيط",
            "bg": "АКТИВИРАНЕ",
            "de": "AKTIVIERUNG",
            "ru": "АКТИВАЦИЯ",
            "uk": "АКТИВАЦІЯ",
            "fr": "ACTIVATION",
            "es": "ACTIVACIÓN",
            "it": "ATTIVAZIONE",
            "sr": "АКТИВАЦИЈА",
            "hr": "AKTIVACIJA",
            "en": "ACTIVATION",
        },
        "hw_a": {
            "ar": "واجب أ — البنية",
            "bg": "ДОМАШНА РАБОТА A — СТРУКТУРА",
            "de": "HAUSAUFGABE A \u2014 STRUKTUR",
            "ru": "ДОМАШНЕЕ ЗАДАНИЕ A \u2014 СТРУКТУРА",
            "uk": "ДОМАШНЄ ЗАВДАННЯ A \u2014 СТРУКТУРА",
            "fr": "DEVOIRS A \u2014 STRUCTURE",
            "es": "TAREA A \u2014 ESTRUCTURA",
            "it": "COMPITI A \u2014 STRUTTURA",
            "sr": "ДОМАЋИ A \u2014 СТРУКТУРА",
            "hr": "ZADAĆA A \u2014 STRUKTURA",
            "en": "HOMEWORK A \u2014 STRUCTURE",
        },
        "hw_b": {
            "ar": "واجب ب — الدقة",
            "bg": "ДОМАШНА РАБОТА B — ПРЕЦИЗНОСТ",
            "de": "HAUSAUFGABE B \u2014 PRÄZISION",
            "ru": "ДОМАШНЕЕ ЗАДАНИЕ B \u2014 ТОЧНОСТЬ",
            "uk": "ДОМАШНЄ ЗАВДАННЯ B \u2014 ТОЧНІСТЬ",
            "fr": "DEVOIRS B \u2014 PRÉCISION",
            "es": "TAREA B \u2014 PRECISIÓN",
            "it": "COMPITI B \u2014 PRECISIONE",
            "sr": "ДОМАЋИ B \u2014 ПРЕЦИЗНОСТ",
            "hr": "ZADAĆA B \u2014 PRECIZNOST",
            "en": "HOMEWORK B \u2014 PRECISION",
        },
        "hw_c": {
            "ar": "واجب ج — طلاقة الكلام",
            "bg": "ДОМАШНА РАБОТА C — ПЛАВНОСТ НА РЕЧТА",
            "de": "HAUSAUFGABE C \u2014 SPRECHFLUSS",
            "ru": "ДОМАШНЕЕ ЗАДАНИЕ C \u2014 БЕГЛОСТЬ РЕЧИ",
            "uk": "ДОМАШНЄ ЗАВДАННЯ C \u2014 ПЛАВНІСТЬ МОВЛЕННЯ",
            "fr": "DEVOIRS C \u2014 FLUIDITÉ ORALE",
            "es": "TAREA C \u2014 FLUIDEZ ORAL",
            "it": "COMPITI C \u2014 FLUENZA ORALE",
            "sr": "ДОМАЋИ C \u2014 ТЕЧНОСТ ГОВОРА",
            "hr": "ZADAĆA C \u2014 TEČNOST GOVORA",
            "en": "HOMEWORK C \u2014 SPEAKING FLUENCY",
        },
        "weekly": {
            "ar": "الخطة الأسبوعية",
            "bg": "СЕДМИЧЕН ПЛАН",
            "de": "WOCHENPLAN",
            "ru": "НЕДЕЛЬНЫЙ ПЛАН",
            "uk": "ТИЖНЕВИЙ ПЛАН",
            "fr": "PLAN HEBDOMADAIRE",
            "es": "PLAN SEMANAL",
            "it": "PIANO SETTIMANALE",
            "sr": "НЕДЕЉНИ ПЛАН",
            "hr": "TJEDNI PLAN",
            "en": "WEEKLY PLAN",
        },
        "mini": {
            "ar": "كتلة صغيرة",
            "bg": "МИНИ-БЛОК",
            "de": "MINI-BLOCK",
            "ru": "МИНИ-БЛОК",
            "uk": "МІНІ-БЛОК",
            "fr": "MINI-BLOC",
            "es": "MINI-BLOQUE",
            "it": "MINI-BLOCCO",
            "sr": "МИНИ-БЛОК",
            "hr": "MINI-BLOK",
            "en": "MINI BLOCK",
        },
        "cefr": {
            "ar": "تقييم CEFR",
            "bg": "ОЦЕНКА CEFR",
            "de": "CEFR-EINSTUFUNG",
            "ru": "ОЦЕНКА CEFR",
            "uk": "ОЦІНКА CEFR",
            "fr": "ÉVALUATION CECR",
            "es": "EVALUACIÓN MCER",
            "it": "VALUTAZIONE QCER",
            "sr": "ЦЕФР ОЦЕНА",
            "hr": "CEFR PROCJENA",
            "en": "CEFR RATING",
        },
    }

    def add_section_banner(key, count=None):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_shading(cell, ACCENT_BG)
        set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        banner = BANNERS.get(key, {})
        assessed_text = banner.get(assessed_code, banner.get("de", key.upper()))
        native_text = banner.get(lang_code, banner.get("en", ""))
        if count:
            assessed_text = assessed_text.replace("{n}", str(count))
            native_text = native_text.replace("{n}", str(count))
        if assessed_code == lang_code:
            display = assessed_text
        elif native_text and native_text != assessed_text:
            display = f"{assessed_text} / {native_text}"
        else:
            display = assessed_text
        run = p.add_run(display)
        run.font.name = 'Arial'; run.font.size = Pt(8); run.font.color.rgb = WHITE; run.font.bold = True
        set_table_borders(tbl, color=ACCENT_BG)

    # ══════════════════════════════════════
    # HELPER: DE/Native column header row
    # ══════════════════════════════════════
    ACCENT_DK = "1A4A7A"
    en_name, native_name = translate_term(native_lang, lang_code)
    # Right column shows just the native script name
    if lang_code == "en":
        native_label = "English"
    else:
        native_label = native_name  # e.g. "Русский", "Français", "Italiano"

    def add_column_headers():
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, label in enumerate([assessed_en, native_label]):
            cell = tbl.cell(0, idx)
            bg = ACCENT_BG if idx == 0 else ACCENT_DK
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(label)
            run.font.name = 'Arial'; run.font.size = Pt(9); run.font.color.rgb = WHITE; run.font.bold = True
        set_table_borders(tbl, color=ACCENT_BG)

    # ══════════════════════════════════════
    # HELPER: "Marko → Student" addressing
    # ══════════════════════════════════════
    def add_addressing():
        marko_native_map = {
            "ru": "Марко", "uk": "Марко", "sr": "Марко", "hr": "Marko",
            "fr": "Marko", "es": "Marko", "it": "Marko", "en": "Marko",
        }
        marko_n = marko_native_map.get(lang_code, "Marko")
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Left: German
        cell_l = tbl.cell(0, 0)
        set_cell_margins(cell_l, top=40, bottom=40, left=80, right=80)
        p = cell_l.paragraphs[0]
        run = p.add_run(f"Marko \u2192 {display}")
        run.font.name = 'Arial'; run.font.size = Pt(9); run.font.color.rgb = INK; run.font.bold = True
        # Right: Native
        cell_r = tbl.cell(0, 1)
        set_cell_margins(cell_r, top=40, bottom=40, left=80, right=80)
        p2 = cell_r.paragraphs[0]
        run2 = p2.add_run(f"{marko_n} \u2192 {native_display}")
        run2.font.name = 'Arial'; run2.font.size = Pt(9); run2.font.color.rgb = INK; run2.font.bold = True
        set_table_borders(tbl)

    # ══════════════════════════════════════
    # HELPER: Side-by-side content row
    # ══════════════════════════════════════
    def add_side_by_side(text_de, text_native, bold=False, color_de=INK, color_native=INK_SOFT, size=9, bg=None):
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, (text, color) in enumerate([(text_de, color_de), (text_native, color_native)]):
            cell = tbl.cell(0, idx)
            if bg:
                set_cell_shading(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
            p = cell.paragraphs[0]
            # Handle multi-line
            lines = text.split('\n')
            for j, line in enumerate(lines):
                if j > 0:
                    p = cell.add_paragraph()
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                run = p.add_run(line)
                run.font.name = 'Arial'
                run.font.size = Pt(size)
                run.font.color.rgb = color
                run.font.bold = bold
        set_table_borders(tbl)

    def add_spacer():
        p = doc.add_paragraph("")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        pf = p.paragraph_format
        pf.line_spacing = Pt(4)

    # ══════════════════════════════════════
    # SECTION: Perception + Strengths + Core Insight
    # ══════════════════════════════════════
    add_column_headers()
    add_addressing()

    sec = sections.get("perception")
    if sec:
        add_side_by_side(
            f"\n{sec['title_de']}\n\n{sec['content_de']}",
            f"\n{sec['title_native']}\n\n{sec['content_native']}",
            size=9
        )

    # ══════════════════════════════════════
    # SECTION: Strengths (Full only)
    # ══════════════════════════════════════
    sec = sections.get("strengths")
    if sec and package == "full":
        # Strengths inline
        de_text = f"\n{sec['title_de']}\n"
        native_text = f"\n{sec['title_native']}\n"
        assessed_strength = get_native_label("Was du gut machst", assessed_code)
        if assessed_code == "de": assessed_strength = "Was du gut machst"
        nl_strength = get_native_label("Was du gut machst", lang_code)
        for item in sec.get("items", []):
            de_text += f"\n{assessed_strength}: {item['strength_de']}\n{item['example_de']}\n"
            native_text += f"\n{nl_strength}: {item['strength_native']}\n{item['example_native']}\n"
        add_side_by_side(de_text, native_text, size=9)

    # ══════════════════════════════════════
    # SECTION: Core Insight
    # ══════════════════════════════════════
    sec = sections.get("core_insight")
    if sec:
        add_side_by_side(
            f"\n{sec['title_de']}\n\n{sec['content_de']}",
            f"\n{sec['title_native']}\n\n{sec['content_native']}",
            size=9
        )
    add_spacer()

    # ══════════════════════════════════════
    # SECTION: Problems
    # ══════════════════════════════════════
    sec = sections.get("problems")
    if sec:
        count = len(sec.get("items", []))
        add_section_banner("problems", count=count)
        for i, item in enumerate(sec.get("items", [])):
            add_side_by_side(
                f"{i+1}. {item['name_de']} [{item.get('error_class', '')}]",
                f"{i+1}. {item['name_native']} [{get_native_label(item.get('error_class', ''), lang_code)}]",
                bold=True, color_de=ERROR_C, color_native=ERROR_C, size=9
            )
            add_side_by_side(
                item["explanation_de"],
                item["explanation_native"],
                size=9
            )
            if item.get("evidence_de") and package == "full":
                add_side_by_side(
                    item["evidence_de"],
                    item["evidence_native"],
                    color_de=INK_LIGHT, color_native=INK_LIGHT, size=8, bg=LIGHT_BG
                )
        add_spacer()

    # ══════════════════════════════════════
    # SECTION: Corrections
    # ══════════════════════════════════════
    sec = sections.get("corrections")
    if sec:
        count = len(sec.get("items", []))
        add_section_banner("corrections", count=count)
        for item in sec.get("items", []):
            add_side_by_side(
                item["original_de"],
                item["original_native"],
                color_de=ERROR_C, color_native=ERROR_C, size=9
            )
            add_side_by_side(
                item["corrected_de"],
                item["corrected_native"],
                color_de=SUCCESS, color_native=SUCCESS, size=9
            )
            add_side_by_side(
                item["explanation_de"],
                item["explanation_native"],
                size=8, color_de=INK_SOFT, color_native=INK_LIGHT
            )
            add_spacer()
        add_spacer()

    # ══════════════════════════════════════
    # SECTION: Solutions
    # ══════════════════════════════════════
    sec = sections.get("solutions")
    if sec:
        count = len(sec.get("items", []))
        add_section_banner("solutions", count=count)
        for i, item in enumerate(sec.get("items", [])):
            add_side_by_side(
                f"{i+1}. {item['name_de']}",
                f"{i+1}. {item['name_native']}",
                bold=True, color_de=ACCENT, color_native=ACCENT, size=9
            )
            add_side_by_side(item["action_de"], item["action_native"], size=9)
            add_side_by_side(
                item["result_de"], item["result_native"],
                size=8, color_de=SUCCESS, color_native=SUCCESS, bg=LIGHT_BG
            )
        add_spacer()

    # ══════════════════════════════════════
    # SECTION: Sentence Engine
    # ══════════════════════════════════════
    sec = sections.get("sentence_engine")
    if sec:
        add_section_banner("engine")
        add_side_by_side(
            sec["name_de"], sec["name_native"],
            bold=True, color_de=ACCENT, color_native=ACCENT, size=10
        )
        if sec.get("explanation_de"):
            add_side_by_side(sec["explanation_de"], sec["explanation_native"], size=9)
        for pat in sec.get("patterns", []):
            add_side_by_side(pat["de"], pat["native"], size=9, bg=LIGHT_BG)
        if sec.get("daily_instruction_de"):
            add_side_by_side(
                sec["daily_instruction_de"], sec["daily_instruction_native"],
                bold=True, size=9
            )
        add_spacer()

    # ══════════════════════════════════════
    # SECTION: Activation
    # ══════════════════════════════════════
    sec = sections.get("activation")
    if sec:
        add_section_banner("activation")
        for s in sec.get("sentences", []):
            add_side_by_side(s["de"], s["native"], size=9)
        add_spacer()

    # ══════════════════════════════════════
    # SECTION: Homework A
    # ══════════════════════════════════════
    sec = sections.get("homework_a")
    if sec:
        add_section_banner("hw_a")
        for i, s in enumerate(sec.get("sentences", [])):
            add_side_by_side(f"{i+1}. {s['de']}", f"{i+1}. {s['native']}", size=9)
        add_spacer()

    # ══════════════════════════════════════
    # SECTION: Homework B
    # ══════════════════════════════════════
    sec = sections.get("homework_b")
    if sec:
        add_section_banner("hw_b")
        for item in sec.get("items", []):
            add_side_by_side(
                item["wrong_de"], item["wrong_native"],
                color_de=ERROR_C, color_native=ERROR_C, size=9
            )
            add_side_by_side(
                item["right_de"], item["right_native"],
                color_de=SUCCESS, color_native=SUCCESS, size=9
            )
        add_spacer()

    # ══════════════════════════════════════
    # SECTION: Homework C (Full only)
    # ══════════════════════════════════════
    sec = sections.get("homework_c")
    if sec and package == "full":
        add_section_banner("hw_c")
        for i, s in enumerate(sec.get("sentences", [])):
            add_side_by_side(f"{i+1}. {s['de']}", f"{i+1}. {s['native']}", size=9)
        add_spacer()

    # ══════════════════════════════════════
    # SECTION: Weekly Plan
    # ══════════════════════════════════════
    sec = sections.get("weekly_plan")
    if sec:
        add_section_banner("weekly")
        for week in sec.get("weeks", []):
            add_side_by_side(
                week["label_de"], week["label_native"],
                bold=True, color_de=ACCENT, color_native=ACCENT, size=10
            )
            assessed_focus = get_native_label("Fokus", assessed_code) if assessed_code != "de" else "Fokus"
            assessed_goal = get_native_label("Ziel", assessed_code) if assessed_code != "de" else "Ziel"
            nl_focus = get_native_label("Fokus", lang_code)
            nl_goal = get_native_label("Ziel", lang_code)
            add_side_by_side(
                f"{assessed_focus}: {week['focus_de']}", f"{nl_focus}: {week['focus_native']}", size=9
            )
            add_side_by_side(
                f"{assessed_goal}: {week['goal_de']}", f"{nl_goal}: {week['goal_native']}",
                size=9, color_de=INK_SOFT, color_native=INK_LIGHT
            )
            if week.get("daily_de"):
                add_side_by_side(
                    week["daily_de"], week["daily_native"],
                    size=8, bg=LIGHT_BG
                )
        add_spacer()

    # ══════════════════════════════════════
    # SECTION: Mini Block (Full only)
    # ══════════════════════════════════════
    sec = sections.get("mini_block")
    if sec and package == "full":
        add_section_banner("mini")
        assessed_goal_lbl = get_native_label("Ziel", assessed_code) if assessed_code != "de" else "Ziel"
        assessed_step_lbl = get_native_label("Schritt", assessed_code) if assessed_code != "de" else "Schritt"
        nl_goal = get_native_label("Ziel", lang_code)
        nl_step = get_native_label("Schritt", lang_code)
        add_side_by_side(
            f"{assessed_goal_lbl}: {sec['goal_de']}", f"{nl_goal}: {sec['goal_native']}",
            bold=True, size=9
        )
        add_side_by_side(
            f"{assessed_step_lbl} 1: {sec['step1_de']}", f"{nl_step} 1: {sec['step1_native']}", size=9
        )
        add_side_by_side(
            f"{assessed_step_lbl} 2: {sec['step2_de']}", f"{nl_step} 2: {sec['step2_native']}", size=9
        )
        add_spacer()

    # ══════════════════════════════════════
    # CEFR FOOTER
    # ══════════════════════════════════════
    if cefr:
        add_section_banner("cefr")
        for key_de, key_native in [
            ("active_de", "active_native"),
            ("passive_de", "passive_native"),
            ("confidence_de", "confidence_native"),
        ]:
            if cefr.get(key_de):
                add_side_by_side(cefr[key_de], cefr[key_native], bold=True, size=10)

        if cefr.get("summary_de"):
            add_side_by_side(cefr["summary_de"], cefr["summary_native"], size=9)

        if package == "full":
            for key_de, key_native in [
                ("stability_de", "stability_native"),
                ("self_correction_de", "self_correction_native"),
                ("reassessment_de", "reassessment_native"),
            ]:
                if cefr.get(key_de):
                    add_side_by_side(cefr[key_de], cefr[key_native], size=9, bg=LIGHT_BG)

            # Re-assessment date with booking link
            if reassess_date_str and reassess_cal_url:
                reassess_booking_labels = {
                    "de": "Nächster Termin buchen",
                    "en": "Book your next session",
                    "fr": "Réserver votre prochaine séance",
                    "es": "Reservar tu próxima sesión",
                    "it": "Prenota la tua prossima sessione",
                    "pt": "Agendar a sua próxima sessão",
                    "ru": "Забронировать следующую сессию",
                    "sr": "Закажите следећу сесију",
                    "hr": "Rezervirajte sljedeću sesiju",
                    "ro": "Rezerva urmatoarea sesiune",
                    "pl": "Zarezerwuj następną sesję",
                    "bg": "Запазете следващата сесия",
                }
                assessed_lbl = reassess_booking_labels.get(assessed_code, reassess_booking_labels["en"])
                native_lbl = reassess_booking_labels.get(lang_code, reassess_booking_labels["en"])

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)

                if assessed_code == lang_code:
                    run = p.add_run(f"{assessed_lbl}: {reassess_date_str}  ")
                else:
                    run = p.add_run(f"{assessed_lbl} / {native_lbl}: {reassess_date_str}  ")
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                run.font.color.rgb = ACCENT
                run.font.bold = True

                add_hyperlink(p, reassess_cal_url, "→ Book Now", color=ACCENT, bold=True, size=9)
        add_spacer()

    # ══════════════════════════════════════
    # BLOCK D — Native Language Summary
    # ══════════════════════════════════════
    if block_d:
        doc.add_page_break()

        # Full-width header for Block D
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_shading(cell, HEADER_BG)
        set_cell_margins(cell, top=120, bottom=120, left=200, right=200)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(block_d.get("title_native", "Summary"))
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.color.rgb = WHITE
        run.font.bold = True

        add_spacer()

        # CEFR summary
        if block_d.get("cefr_summary_native"):
            p = doc.add_paragraph()
            run = p.add_run(block_d["cefr_summary_native"])
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.color.rgb = INK
            p.paragraph_format.space_after = Pt(8)

        # Top problems
        if block_d.get("top_problems_native"):
            p = doc.add_paragraph()
            lbl = block_d.get("label_problems") or get_native_label("Top Problems", lang_code)
            run = p.add_run(lbl)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.color.rgb = ERROR_C
            run.font.bold = True
            p.paragraph_format.space_after = Pt(4)

            for prob in block_d["top_problems_native"]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(f"  {prob}")
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                run.font.color.rgb = INK

        add_spacer()

        # Key recommendation
        if block_d.get("key_recommendation_native"):
            p = doc.add_paragraph()
            lbl = block_d.get("label_recommendation") or get_native_label("Key Recommendation", lang_code)
            run = p.add_run(lbl)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.color.rgb = ACCENT
            run.font.bold = True
            p.paragraph_format.space_after = Pt(4)

            p = doc.add_paragraph()
            run = p.add_run(block_d["key_recommendation_native"])
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.color.rgb = INK
            p.paragraph_format.space_after = Pt(8)

        # Next steps
        if block_d.get("next_steps_native"):
            p = doc.add_paragraph()
            lbl = block_d.get("label_next_steps") or get_native_label("Next Steps", lang_code)
            run = p.add_run(lbl)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.color.rgb = SUCCESS
            run.font.bold = True
            p.paragraph_format.space_after = Pt(4)

            p = doc.add_paragraph()
            run = p.add_run(block_d["next_steps_native"])
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.color.rgb = INK

        # Re-assessment booking CTA in Block D
        if reassess_date_str and reassess_cal_url:
            reassess_cta_labels = {
                "de": "Dein nächstes Assessment",
                "en": "Your Next Assessment",
                "fr": "Votre prochaine évaluation",
                "es": "Tu próxima evaluación",
                "it": "La tua prossima valutazione",
                "pt": "A sua próxima avaliação",
                "ru": "Ваша следующая оценка",
                "sr": "Ваша следећа процена",
                "hr": "Vaša sljedeća procjena",
                "ro": "Următoarea dvs. evaluare",
                "pl": "Twoja następna ocena",
                "bg": "Следващата ви оценка",
                "ar": "تقييمك القادم",
            }
            reassess_price_labels = {
                "de": "Re-Assessment zum Sonderpreis",
                "en": "Re-Assessment at special rate",
                "fr": "Réévaluation au tarif spécial",
                "es": "Reevaluación a tarifa especial",
                "it": "Rivalutazione a tariffa speciale",
                "pt": "Reavaliação a tarifa especial",
                "ru": "Повторная оценка по специальной цене",
                "sr": "Поновна процена по специјалној цени",
                "hr": "Ponovna procjena po posebnoj cijeni",
                "ro": "Reevaluare la tarif special",
                "pl": "Ponowna ocena w specjalnej cenie",
                "bg": "Преоценка на специална цена",
                "ar": "إعادة تقييم بسعر خاص",
            }
            cta_lbl = reassess_cta_labels.get(lang_code, reassess_cta_labels["en"])
            price_lbl = reassess_price_labels.get(lang_code, reassess_price_labels["en"])

            doc.add_paragraph("")
            # Booking box
            book_tbl = doc.add_table(rows=1, cols=1)
            book_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            bc = book_tbl.cell(0, 0)
            set_cell_shading(bc, "EFF6FF")
            set_cell_margins(bc, top=140, bottom=140, left=200, right=200)
            # Set border
            from docx.oxml.ns import qn as _qn
            tc_pr = bc._tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="8" w:space="0" w:color="2563AB"/>'
                f'<w:left w:val="single" w:sz="8" w:space="0" w:color="2563AB"/>'
                f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="2563AB"/>'
                f'<w:right w:val="single" w:sz="8" w:space="0" w:color="2563AB"/>'
                f'</w:tcBorders>'
            )
            tc_pr.append(borders)

            bp = bc.paragraphs[0]
            bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = bp.add_run(cta_lbl)
            run.font.name = 'Arial'; run.font.size = Pt(12); run.font.color.rgb = ACCENT; run.font.bold = True

            bp2 = bc.add_paragraph()
            bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp2.paragraph_format.space_before = Pt(6)
            run = bp2.add_run(reassess_date_str)
            run.font.name = 'Arial'; run.font.size = Pt(11); run.font.color.rgb = INK; run.font.bold = True

            bp3 = bc.add_paragraph()
            bp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp3.paragraph_format.space_before = Pt(4)
            run = bp3.add_run(f"{price_lbl}: ")
            run.font.name = 'Arial'; run.font.size = Pt(9); run.font.color.rgb = INK_SOFT
            run2 = bp3.add_run("€69")
            run2.font.name = 'Arial'; run2.font.size = Pt(11); run2.font.color.rgb = SUCCESS; run2.font.bold = True

            bp4 = bc.add_paragraph()
            bp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp4.paragraph_format.space_before = Pt(8)
            add_hyperlink(bp4, reassess_cal_url, "→ Book Re-Assessment", color=ACCENT, bold=True, size=10)

        # ── Professional closing — Chris Voss: create forward commitment ──
        doc.add_paragraph("")
        close_table = doc.add_table(rows=1, cols=1)
        close_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cc = close_table.cell(0, 0)
        set_cell_shading(cc, HEADER_BG)
        set_cell_margins(cc, top=120, bottom=120, left=200, right=200)

        cp = cc.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run("lingo")
        run.font.name = 'Arial'; run.font.size = Pt(12); run.font.color.rgb = WHITE; run.font.bold = True
        run1b = cp.add_run("grade")
        run1b.font.name = 'Arial'; run1b.font.size = Pt(12); run1b.font.color.rgb = RGBColor(0x93, 0xB5, 0xD0); run1b.font.bold = True
        run1c = cp.add_run(".com")
        run1c.font.name = 'Arial'; run1c.font.size = Pt(12); run1c.font.color.rgb = RGBColor(0x93, 0xB5, 0xD0); run1c.font.bold = True

        cp2 = cc.add_paragraph(); cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = cp2.add_run("Professional Language Assessment")
        run2.font.name = 'Arial'; run2.font.size = Pt(9); run2.font.color.rgb = RGBColor(0x93, 0xB5, 0xD0)

        cp3 = cc.add_paragraph(); cp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Translate closing line
        closing_map = {
            "ru": f"Этот отчёт подготовлен Марко для {native_display}.",
            "uk": f"Цей звіт підготовлений Марко для {native_display}.",
            "fr": f"Ce rapport a été préparé par Marko pour {native_display}.",
            "es": f"Este informe fue preparado por Marko para {native_display}.",
            "it": f"Questo rapporto è stato preparato da Marko per {native_display}.",
            "sr": f"Овај извештај је припремио Марко за {native_display}.",
            "hr": f"Ovaj izvještaj je pripremio Marko za {native_display}.",
            "bg": f"Този доклад е подготвен от Марко за {native_display}.",
            "ar": f"أعد هذا التقرير ماركو لـ {native_display}.",
            "pl": f"Ten raport przygotował Marko dla {native_display}.",
            "tr": f"Bu rapor Marko tarafından {native_display} için hazırlanmıştır.",
            "sq": f"Ky raport u përgatit nga Marko për {native_display}.",
            "ro": f"Acest raport a fost pregătit de Marko pentru {native_display}.",
            "hu": f"Ezt a jelentést Marko készítette {native_display} számára.",
        }
        closing_text = closing_map.get(lang_code, f"This report was prepared by Marko for {display}.")
        run3 = cp3.add_run(closing_text)
        run3.font.name = 'Arial'; run3.font.size = Pt(8); run3.font.color.rgb = RGBColor(0x93, 0xB5, 0xD0)
        run3.font.italic = True

    # ══════════════════════════════════════
    # DOCUMENT FOOTER on every page
    # ══════════════════════════════════════
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("lingo")
    run.font.name = 'Arial'; run.font.size = Pt(7); run.font.color.rgb = INK
    run.font.bold = True
    run1b = fp.add_run("grade")
    run1b.font.name = 'Arial'; run1b.font.size = Pt(7); run1b.font.color.rgb = RGBColor(0x93, 0xB5, 0xD0)
    run1b.font.bold = True
    run1c = fp.add_run(".com  |  Confidential Language Assessment Report")
    run1c.font.name = 'Arial'; run1c.font.size = Pt(7); run1c.font.color.rgb = INK_LIGHT

    # ══════════════════════════════════════
    # DOCUMENT HEADER on every page
    # ══════════════════════════════════════
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if HAS_LOGO:
        run = hp.add_run()
        run.add_picture(LOGO_PATH, height=Cm(0.7))
        run2 = hp.add_run("  Lingo")
        run2.font.name = 'Arial'
        run2.font.size = Pt(8)
        run2.font.color.rgb = NAVY
        run2.font.bold = True
        run2b = hp.add_run("Grade")
        run2b.font.name = 'Arial'
        run2b.font.size = Pt(8)
        run2b.font.color.rgb = RGBColor(0x93, 0xB5, 0xD0)
        run2b.font.bold = True
        run3 = hp.add_run("  |  Professional Language Assessment")
        run3.font.name = 'Arial'
        run3.font.size = Pt(8)
        run3.font.color.rgb = INK_LIGHT
    else:
        run = hp.add_run("Lingo")
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        run.font.color.rgb = NAVY
        run.font.bold = True
        run1b = hp.add_run("Grade")
        run1b.font.name = 'Arial'
        run1b.font.size = Pt(8)
        run1b.font.color.rgb = RGBColor(0x93, 0xB5, 0xD0)
        run1b.font.bold = True
        run2 = hp.add_run("  |  Professional Language Assessment")
        run2.font.name = 'Arial'
        run2.font.size = Pt(8)
        run2.font.color.rgb = INK_LIGHT

    # Save
    doc.save(output_path)
    return output_path


# ── Sample JSON for testing ──
SAMPLE_JSON = {
    "metadata": {
        "student_name": "Viktor K.",
        "date": "2026-03-19",
        "package": "full",
        "assessed_language": "Deutsch",
        "native_language": "Italienisch",
        "assessor": "Marko",
        "active_cefr": "B1.2",
        "passive_cefr": "B2.1",
        "confidence": "78%",
        "structural_stability": "mittel",
        "self_correction": "mittel",
        "sentence_control": "teilstabil",
        "recommended_reassessment": "3 Monate",
        "error_focus": ["Verbposition", "Kasus/Artikel"]
    },
    "sections": [
        {
            "id": "perception",
            "title_de": "Präzise Wahrnehmung",
            "title_native": "Precise Perception",
            "content_de": "Du hast in dieser Sitzung gezeigt, dass du ein solides Grundgerüst im Deutschen hast. Dein Wortschatz ist angemessen für B1-Kommunikation, und du hast dich nicht gescheut, komplexere Strukturen zu versuchen. Gleichzeitig sind unter Sprechdruck deutliche Muster sichtbar geworden: Nebensatzstrukturen brechen zusammen, das Kasussystem ist instabil, und die Verbposition im Nebensatz wird regelmäßig verfehlt.",
            "content_native": "In this session you showed that you have a solid foundation in German. Your vocabulary is adequate for B1 communication, and you weren't afraid to attempt more complex structures. At the same time, clear patterns became visible under speaking pressure: subordinate clause structures break down, the case system is unstable, and verb position in subordinate clauses is regularly missed."
        },
        {
            "id": "strengths",
            "title_de": "Strukturelle Stärken",
            "title_native": "Structural Strengths",
            "items": [
                {
                    "strength_de": "Solide Verbzweitstellung in Hauptsätzen",
                    "strength_native": "Solid verb-second placement in main clauses",
                    "example_de": "\"Gestern bin ich ins Kino gegangen.\" — Korrekte Inversion nach temporalem Adverb.",
                    "example_native": "\"Yesterday I went to the cinema.\" — Correct inversion after temporal adverb."
                },
                {
                    "strength_de": "Guter Einsatz von Modalverben",
                    "strength_native": "Good use of modal verbs",
                    "example_de": "\"Ich muss morgen früher aufstehen.\" — Korrekte Klammerstruktur.",
                    "example_native": "\"I have to get up earlier tomorrow.\" — Correct bracket structure."
                },
                {
                    "strength_de": "Aktive Selbstkorrektur-Versuche",
                    "strength_native": "Active self-correction attempts",
                    "example_de": "Du hast mehrfach angesetzt, Fehler zu korrigieren — auch wenn die Korrektur nicht immer gelang, zeigt es Sprachbewusstsein.",
                    "example_native": "You made multiple attempts to correct errors — even when the correction didn't always succeed, it shows language awäreness."
                }
            ]
        },
        {
            "id": "core_insight",
            "title_de": "Zentrale Struktur-Erkenntnis",
            "title_native": "Core Structural Insight",
            "content_de": "Das zentrale Muster dieser Sitzung ist die Instabilität der Nebensatzstruktur unter Druck. In ruhigen Momenten gelingt dir die korrekte Verbendstellung. Sobald der Sprechdruck steigt, fällst du in die Hauptsatz-Wortstellung zurück. Der Motor dagegen ist systematisches Drilling der weil/dass/obwohl-Struktur.",
            "content_native": "The central pattern of this session is the instability of subordinate clause structure under pressure. In calm moments you achieve correct verb-final position. As soon as speaking pressure increases, you fall back into main clause word order. The engine against this is systematic drilling of because/that/although structures."
        },
        {
            "id": "problems",
            "title_de": "Top Strukturelle Probleme",
            "title_native": "Top Structural Problems",
            "items": [
                {
                    "name_de": "Verbposition im Nebensatz",
                    "name_native": "Verb position in subordinate clauses",
                    "error_class": "Verbposition",
                    "explanation_de": "Du hast das Verb im Nebensatz regelmäßig in die Zweitposition gestellt, statt es ans Ende zu setzen. Das ist der häufigste Strukturfehler auf B1-Niveau.",
                    "explanation_native": "You regularly placed the verb in second position in subordinate clauses instead of putting it at the end. This is the most common structural error at B1 level.",
                    "evidence_de": "\"...weil ich habe keine Zeit.\" statt \"...weil ich keine Zeit habe.\"",
                    "evidence_native": "\"...because I have no time.\" instead of \"...because I no time have.\""
                },
                {
                    "name_de": "Akkusativ/Dativ-Verwechslung",
                    "name_native": "Accusative/Dative confusion",
                    "error_class": "Kasus/Artikel",
                    "explanation_de": "Wechselpräpositionen wurden durchgehend mit Akkusativ verwendet, auch in statischen Kontexten, wo Dativ korrekt wäre.",
                    "explanation_native": "Two-way prepositions were consistently used with accusative, even in static contexts where dative would be correct.",
                    "evidence_de": "\"Ich bin in den Park.\" statt \"Ich bin im Park.\"",
                    "evidence_native": "\"I am in the park (motion).\" instead of \"I am in the park (location).\""
                },
                {
                    "name_de": "Perfekt-Bildung mit Bewegungsverben",
                    "name_native": "Perfect tense with motion verbs",
                    "error_class": "Zeitformen",
                    "explanation_de": "Bewegungsverben wurden mit 'haben' statt 'sein' konjugiert. Das betrifft gehen, fahren, kommen, fliegen.",
                    "explanation_native": "Motion verbs were conjugated with 'have' instead of 'be'. This affects go, drive, come, fly.",
                    "evidence_de": "\"Ich habe gegangen.\" statt \"Ich bin gegangen.\"",
                    "evidence_native": "\"I have went.\" instead of \"I have gone.\""
                },
                {
                    "name_de": "Fehlende Satzklammer bei trennbaren Verben",
                    "name_native": "Missing sentence bracket with separable verbs",
                    "error_class": "Verbposition",
                    "explanation_de": "Trennbare Verben wurden nicht getrennt. Das Präfix blieb am Verb kleben.",
                    "explanation_native": "Separable verbs were not separated. The prefix stayed attached to the verb.",
                    "evidence_de": "\"Ich aufstehe um 7 Uhr.\" statt \"Ich stehe um 7 Uhr auf.\"",
                    "evidence_native": "\"I getup at 7.\" instead of \"I get up at 7.\""
                },
                {
                    "name_de": "Wortstellung nach temporalen Adverbien",
                    "name_native": "Word order after temporal adverbs",
                    "error_class": "Wortstellung",
                    "explanation_de": "Nach Zeitangaben am Satzanfang wurde das Subjekt nicht hinter das Verb gestellt.",
                    "explanation_native": "After time expressions at the beginning of sentences, the subject was not placed after the verb.",
                    "evidence_de": "\"Gestern ich bin ins Kino gegangen.\" statt \"Gestern bin ich ins Kino gegangen.\"",
                    "evidence_native": "\"Yesterday I am to cinema went.\" instead of \"Yesterday went I to the cinema.\""
                }
            ]
        },
        {
            "id": "corrections",
            "title_de": "Hebel-Korrekturen",
            "title_native": "Lever Corrections",
            "items": [
                {
                    "original_de": "Was du gesagt hast: \"Ich habe gestern gegeht ins Kino.\"",
                    "original_native": "What you said: \"I have yesterday went to cinema.\"",
                    "corrected_de": "Saubere Version: \"Ich bin gestern ins Kino gegangen.\"",
                    "corrected_native": "Clean version: \"I went to the cinema yesterday.\"",
                    "explanation_de": "Bewegungsverben nutzen 'sein' als Hilfsverb, und das Partizip steht am Ende.",
                    "explanation_native": "Motion verbs use 'sein' as auxiliary, and the past participle goes at the end."
                },
                {
                    "original_de": "Was du gesagt hast: \"...weil ich habe keine Zeit.\"",
                    "original_native": "What you said: \"...because I have no time.\"",
                    "corrected_de": "Saubere Version: \"...weil ich keine Zeit habe.\"",
                    "corrected_native": "Clean version: \"...because I no time have.\"",
                    "explanation_de": "Nach 'weil' wandert das konjugierte Verb ans Sätzende. Immer.",
                    "explanation_native": "After 'because' the conjugated verb moves to the end. Always."
                },
                {
                    "original_de": "Was du gesagt hast: \"Ich aufstehe um 7 Uhr.\"",
                    "original_native": "What you said: \"I getup at 7 o'clock.\"",
                    "corrected_de": "Saubere Version: \"Ich stehe um 7 Uhr auf.\"",
                    "corrected_native": "Clean version: \"I get up at 7 o'clock.\"",
                    "explanation_de": "Trennbare Verben: konjugierter Teil an Position 2, Präfix ans Ende.",
                    "explanation_native": "Separable verbs: conjugated part in position 2, prefix to the end."
                },
                {
                    "original_de": "Was du gesagt hast: \"Ich bin in den Park.\"",
                    "original_native": "What you said: \"I am in the park (motion).\"",
                    "corrected_de": "Saubere Version: \"Ich bin im Park.\"",
                    "corrected_native": "Clean version: \"I am in the park (location).\"",
                    "explanation_de": "Statisch = Dativ. 'in dem' wird zu 'im'. Keine Bewegung = kein Akkusativ.",
                    "explanation_native": "Static = Dative. No movement = no accusative."
                },
                {
                    "original_de": "Was du gesagt hast: \"Gestern ich bin ins Kino gegangen.\"",
                    "original_native": "What you said: \"Yesterday I am to cinema went.\"",
                    "corrected_de": "Saubere Version: \"Gestern bin ich ins Kino gegangen.\"",
                    "corrected_native": "Clean version: \"Yesterday I went to the cinema.\"",
                    "explanation_de": "Verb bleibt IMMER an Position 2. Wenn eine Zeitangabe den Satz eröffnet, rutscht das Subjekt hinter das Verb.",
                    "explanation_native": "The verb ALWAYS stays in position 2. When a time expression opens the sentence, the subject slides behind the verb."
                }
            ]
        },
        {
            "id": "solutions",
            "title_de": "Top Lösungen",
            "title_native": "Top Solutions",
            "items": [
                {
                    "name_de": "Satzmotor: weil + S + ... + Verb",
                    "name_native": "Sentence Engine: because + S + ... + Verb",
                    "action_de": "Jeden Tag 5 Sätze mit 'weil' bilden. Laut sprechen. Verb bewusst ans Ende setzen.",
                    "action_native": "Form 5 sentences with 'because' every day. Speak aloud. Consciously place verb at the end.",
                    "result_de": "Nach 2 Wochen: Nebensatzstruktur automatisiert.",
                    "result_native": "After 2 weeks: subordinate clause structure automated."
                },
                {
                    "name_de": "Bewegungsverben-Liste",
                    "name_native": "Motion Verbs List",
                    "action_de": "Liste der 10 wichtigsten Bewegungsverben lernen. Jeden Tag 3 Sätze im Perfekt mit 'sein' bilden.",
                    "action_native": "Learn the list of 10 most important motion verbs. Form 3 perfect tense sentences with 'sein' every day.",
                    "result_de": "Nach 2 Wochen: sein/haben-Auswahl automatisiert.",
                    "result_native": "After 2 weeks: sein/haben selection automated."
                },
                {
                    "name_de": "Dativ-Trigger trainieren",
                    "name_native": "Dative Trigger Training",
                    "action_de": "3 Sätze pro Tag mit Wechselpräpositionen. Jedes Mal fragen: Bewegung oder Position?",
                    "action_native": "3 sentences per day with two-way prepositions. Each time ask: movement or position?",
                    "result_de": "Nach 3 Wochen: Kasus-Entscheidung wird reflexartig.",
                    "result_native": "After 3 weeks: case decision becomes reflexive."
                },
                {
                    "name_de": "Trennbare-Verben-Drill",
                    "name_native": "Separable Verb Drill",
                    "action_de": "5 trennbare Verben pro Woche. Jeden Tag 2 Sätze pro Verb. Präfix bewusst ans Ende.",
                    "action_native": "5 separable verbs per week. 2 sentences per verb every day. Prefix consciously to the end.",
                    "result_de": "Nach 2 Wochen: Satzklammer stabil.",
                    "result_native": "After 2 weeks: sentence bracket stable."
                },
                {
                    "name_de": "Inversions-Training",
                    "name_native": "Inversion Training",
                    "action_de": "Jeden Satz mit einer Zeitangabe beginnen. Sofort Verb, dann Subjekt.",
                    "action_native": "Start every sentence with a time expression. Immediately verb, then subject.",
                    "result_de": "Nach 1 Woche: Inversion wird Gewohnheit.",
                    "result_native": "After 1 week: inversion becomes habit."
                }
            ]
        },
        {
            "id": "sentence_engine",
            "title_de": "Satzmotor",
            "title_native": "Sentence Engine",
            "name_de": "weil + Subjekt + ... + Verb (am Ende)",
            "name_native": "because + Subject + ... + Verb (at end)",
            "explanation_de": "Dieser Motor adressiert dein größtes strukturelles Problem: die Verbposition im Nebensatz. Wenn du diesen Motor automatisierst, lösen sich 60% deiner Strukturfehler von allein.",
            "explanation_native": "This engine addresses your biggest structural problem: verb position in subordinate clauses. When you automate this engine, 60% of your structural errors resolve themselves.",
            "patterns": [
                {"de": "Ich bin müde, weil ich gestern lange gearbeitet habe.", "native": "I am tired because I worked long hours yesterday."},
                {"de": "Er kommt nicht, weil er krank ist.", "native": "He's not coming because he is sick."},
                {"de": "Ich lerne Deutsch, weil ich in Deutschland arbeiten möchte.", "native": "I'm learning German because I want to work in Germany."},
                {"de": "Sie hat angerufen, weil sie eine Frage hatte.", "native": "She called because she had a question."},
                {"de": "Wir bleiben zu Hause, weil es regnet.", "native": "We're staying home because it's raining."}
            ],
            "daily_instruction_de": "Sprich jeden Morgen 5 neue weil-Sätze laut. Betone bewusst das Verb am Ende.",
            "daily_instruction_native": "Speak 5 new because-sentences aloud every morning. Consciously emphasize the verb at the end."
        },
        {
            "id": "activation",
            "title_de": "Aktivierung",
            "title_native": "Activation",
            "sentences": [
                {"de": "Ich bin müde, weil ich lange gearbeitet habe.", "native": "I am tired because I worked for a long time."},
                {"de": "Ich komme später, weil ich noch einen Termin habe.", "native": "I'm coming later because I still have an appointment."},
                {"de": "Er ist glücklich, weil er die Prüfung bestanden hat.", "native": "He is happy because he passed the exam."}
            ]
        },
        {
            "id": "homework_a",
            "title_de": "Hausaufgaben A — Struktur",
            "title_native": "Homework A — Structure",
            "sentences": [
                {"de": "Ich ging ins Büro, weil ich einen Termin hatte.", "native": "I went to the office because I had an appointment."},
                {"de": "Er blieb zu Hause, weil er krank war.", "native": "He stayed home because he was sick."},
                {"de": "Wir fuhren nach Berlin, weil wir Freunde besuchen wollten.", "native": "We drove to Berlin because we wanted to visit friends."},
                {"de": "Sie rief mich an, weil sie Hilfe brauchte.", "native": "She called me because she needed help."},
                {"de": "Ich stand früh auf, weil ich zum Flughafen musste.", "native": "I got up early because I had to go to the airport."},
                {"de": "Er lernte Deutsch, weil er in München arbeiten wollte.", "native": "He learned German because he wanted to work in Munich."}
            ]
        },
        {
            "id": "homework_b",
            "title_de": "Hausaufgaben B — Präzision",
            "title_native": "Homework B — Precision",
            "items": [
                {"wrong_de": "Falsch: ...weil ich habe keine Zeit.", "wrong_native": "Wrong: ...because I have no time.", "right_de": "Richtig: ...weil ich keine Zeit habe.", "right_native": "Correct: ...because I no time have."},
                {"wrong_de": "Falsch: Ich habe gegangen.", "wrong_native": "Wrong: I have went.", "right_de": "Richtig: Ich bin gegangen.", "right_native": "Correct: I have gone."},
                {"wrong_de": "Falsch: Ich bin in den Park. (statisch)", "wrong_native": "Wrong: I am in the park (motion).", "right_de": "Richtig: Ich bin im Park.", "right_native": "Correct: I am in the park."},
                {"wrong_de": "Falsch: Ich aufstehe um 7 Uhr.", "wrong_native": "Wrong: I getup at 7.", "right_de": "Richtig: Ich stehe um 7 Uhr auf.", "right_native": "Correct: I get up at 7."},
                {"wrong_de": "Falsch: Gestern ich bin gegangen.", "wrong_native": "Wrong: Yesterday I am went.", "right_de": "Richtig: Gestern bin ich gegangen.", "right_native": "Correct: Yesterday I went."},
                {"wrong_de": "Falsch: Er hat gefahrt.", "wrong_native": "Wrong: He has drived.", "right_de": "Richtig: Er ist gefahren.", "right_native": "Correct: He drove / has driven."},
                {"wrong_de": "Falsch: Ich möchte in den Kino gehen.", "wrong_native": "Wrong: I want to go in the cinema.", "right_de": "Richtig: Ich möchte ins Kino gehen.", "right_native": "Correct: I want to go to the cinema."},
                {"wrong_de": "Falsch: ...dass er kommt morgen.", "wrong_native": "Wrong: ...that he comes tomorrow.", "right_de": "Richtig: ...dass er morgen kommt.", "right_native": "Correct: ...that he tomorrow comes."}
            ]
        },
        {
            "id": "homework_c",
            "title_de": "Hausaufgaben C — Sprechfluss",
            "title_native": "Homework C — Speaking Fluency",
            "sentences": [
                {"de": "Zuerst machte ich mir einen Kaffee.", "native": "First I made myself a coffee."},
                {"de": "Dann setzte ich mich an den Schreibtisch.", "native": "Then I sat down at the desk."},
                {"de": "Ich öffnete meinen Laptop, weil ich arbeiten musste.", "native": "I opened my laptop because I had to work."},
                {"de": "Danach las ich meine E-Mails.", "native": "After that I read my emails."},
                {"de": "Außerdem rief ich einen Kollegen an.", "native": "Additionally I called a colleague."},
                {"de": "Obwohl ich müde war, arbeitete ich konzentriert.", "native": "Although I was tired, I worked with concentration."},
                {"de": "Weil ich einen Termin hatte, ging ich früh los.", "native": "Because I had an appointment, I left early."},
                {"de": "Bevor ich das Haus verließ, machte ich mir ein Brot.", "native": "Before I left the house, I made myself a sandwich."},
                {"de": "Nachdem ich gegessen hatte, fuhr ich mit dem Bus.", "native": "After I had eaten, I took the bus."},
                {"de": "Im Bus las ich ein Buch, weil die Fahrt lang war.", "native": "On the bus I read a book because the ride was long."},
                {"de": "Schließlich kam ich im Büro an.", "native": "Finally I arrived at the office."},
                {"de": "Es war ein produktiver Morgen, obwohl er anstrengend war.", "native": "It was a productive morning, although it was exhausting."}
            ]
        },
        {
            "id": "weekly_plan",
            "title_de": "Wochenplan",
            "title_native": "Weekly Plan",
            "weeks": [
                {"label_de": "Woche 1-2", "label_native": "Week 1-2", "focus_de": "Nebensatzstruktur mit weil und dass", "focus_native": "Subordinate clause structure with because and that", "goal_de": "Automatisierung der Verbendstellung", "goal_native": "Automation of verb-final position", "daily_de": "5 weil/dass-Sätze pro Tag laut sprechen", "daily_native": "Speak 5 because/that sentences aloud per day"},
                {"label_de": "Woche 3-4", "label_native": "Week 3-4", "focus_de": "Bewegungsverben + sein/haben", "focus_native": "Motion verbs + sein/haben", "goal_de": "Korrekte Hilfsverb-Auswahl", "goal_native": "Correct auxiliary verb selection", "daily_de": "3 Perfekt-Sätze mit Bewegungsverben pro Tag", "daily_native": "3 perfect tense sentences with motion verbs per day"},
                {"label_de": "Woche 5-6", "label_native": "Week 5-6", "focus_de": "Wechselpräpositionen + Kasus", "focus_native": "Two-way prepositions + case", "goal_de": "Bewegung vs. Position reflexartig unterscheiden", "goal_native": "Reflexively distinguish movement vs. position", "daily_de": "3 Sätze mit in/auf/an + Bewegung/Position", "daily_native": "3 sentences with in/on/at + movement/position"},
                {"label_de": "Woche 7-8", "label_native": "Week 7-8", "focus_de": "Trennbare Verben + Inversion", "focus_native": "Separable verbs + Inversion", "goal_de": "Satzklammer und V2-Regel stabilisieren", "goal_native": "Stabilize sentence bracket and V2 rule", "daily_de": "2 trennbare Verben + 3 Inversions-Sätze pro Tag", "daily_native": "2 separable verbs + 3 inversion sentences per day"}
            ]
        },
        {
            "id": "mini_block",
            "title_de": "Mini-Block",
            "title_native": "Mini Block",
            "goal_de": "Nebensätze unter Sprechdruck stabilisieren",
            "goal_native": "Stabilize subordinate clauses under speaking pressure",
            "step1_de": "Taglich 5 weil-Sätze laut sprechen — bewusst langsam, Verb betont am Ende.",
            "step1_native": "Speak 5 because-sentences aloud daily — consciously slow, verb emphasized at the end.",
            "step2_de": "Ein 2-Minuten-Monolog mit mindestens 3 Nebensätzen. Aufnehmen. Anhoren. Wiederholen.",
            "step2_native": "A 2-minute monologue with at least 3 subordinate clauses. Record. Listen. Repeat."
        }
    ],
    "cefr_footer": {
        "active_de": "Aktive gesprochene Produktion: B1.2",
        "active_native": "Active spoken production: B1.2",
        "passive_de": "Passive Kompetenz (geschätzte Verstehensobergrenze): B2.1",
        "passive_native": "Passive competence (estimated comprehension ceiling): B2.1",
        "confidence_de": "Konfidenz: 78%",
        "confidence_native": "Confidence: 78%",
        "summary_de": "Du befindest dich im oberen B1-Bereich. Dein Grundgerüst ist vorhanden, aber unter Druck instabil. Mit gezieltem Strukturtraining ist B2 in 3-4 Monaten realistisch.",
        "summary_native": "You are in the upper B1 range. Your foundation is in place but unstable under pressure. With targeted structural training, B2 is realistic in 3-4 months.",
        "stability_de": "Strukturstabilität unter Druck: mittel",
        "stability_native": "Structural stability under pressure: medium",
        "self_correction_de": "Selbstkorrektur-Fähigkeit: mittel",
        "self_correction_native": "Self-correction ability: medium",
        "reassessment_de": "Empfohlenes Re-Assessment: 3 Monate",
        "reassessment_native": "Recommended re-assessment: 3 months"
    },
    "block_d": {
        "title_native": "Summary in Your Language",
        "cefr_summary_native": "Your active spoken German is at level B1.2. This means you can communicate in everyday situations, but complex structures break down under pressure. Your passive understanding is higher at B2.1.",
        "top_problems_native": [
            "Verb position in subordinate clauses (weil, dass, obwohl)",
            "Accusative/Dative confusion with two-way prepositions",
            "Perfect tense formation with motion verbs (sein vs. haben)",
            "Separable verbs not being separated",
            "Word order after temporal adverbs (inversion)"
        ],
        "key_recommendation_native": "Practice the Sentence Engine daily: form 5 sentences with 'weil' (because) and consciously place the verb at the end. This single drill will fix your biggest structural weakness.",
        "next_steps_native": "Focus on subordinate clause structure for the next 2 weeks (weil + dass), then move to motion verbs and case system. Book a re-assessment in 3 months to measure progress."
    }
}


if __name__ == "__main__":
    if len(sys.argv) > 1:
        input_path = sys.argv[1]
        with open(input_path, 'r', encoding='utf-8-sig') as f:
            data = json.load(f)
        output_path = sys.argv[2] if len(sys.argv) > 2 else None
    else:
        # Use sample data for testing
        data = SAMPLE_JSON
        output_path = "LingoGrade_Report_Sample.docx"

    result = create_report(data, output_path)
    print(f"Report generated: {result}")
